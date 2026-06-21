from docx import Document
import os

doc = Document()
doc.add_heading('CollegeBuddy - Project Overview & Use Cases', 0)

doc.add_heading('Project Overview', level=1)
doc.add_paragraph('CollegeBuddy is a comprehensive, full-stack smart campus event management platform designed to revolutionize how events are discovered, managed, and executed within educational institutions. It streamlines the entire lifecycle of a campus event. It provides a seamless experience for students to find events, hosts to manage them with AI-driven insights, and volunteers to verify attendance using a robust QR-based smart ticketing system.')

doc.add_heading('Technology Stack', level=1)
doc.add_paragraph('Frontend: React 19 (Vite) with TailwindCSS for styling, React Router for navigation, Axios for API communication, and html5-qrcode for in-browser QR scanning.', style='List Bullet')
doc.add_paragraph('Backend: FastAPI (Python 3.10+) serving the REST API.', style='List Bullet')
doc.add_paragraph('Database: PostgreSQL (with SQLite fallback for development) using SQLAlchemy ORM.', style='List Bullet')
doc.add_paragraph('Security: Secure JWT-based authentication and Role-Based Access Control (RBAC), with Passlib (Bcrypt) for password hashing.', style='List Bullet')
doc.add_paragraph('AI/ML engine: Scikit-Learn (Linear Regression model) stored via Joblib to power predictive analytics.', style='List Bullet')

doc.add_heading('User Roles and Workflows', level=1)
doc.add_heading('1. Students (Attendees):', level=2)
doc.add_paragraph('Can browse through a dynamic list of upcoming campus events.', style='List Bullet')
doc.add_paragraph('Register/Book tickets with a single click.', style='List Bullet')
doc.add_paragraph('Receive a unique, secure QR code ticket for every registered event (stored locally and in the database).', style='List Bullet')
doc.add_paragraph('Manage their user profile and event history.', style='List Bullet')

doc.add_heading('2. Hosts (Organizers):', level=2)
doc.add_paragraph('Have access to a dashboard to create, edit, and manage rich events (including setting participant limits, fees, and uploading posters).', style='List Bullet')
doc.add_paragraph('View real-time analytics and statistics on event registrations.', style='List Bullet')
doc.add_paragraph('Leverage an AI prediction model to estimate actual event attendance based on metrics like participant limits, fees, and current registration numbers.', style='List Bullet')

doc.add_heading('3. Volunteers (Event Staff):', level=2)
doc.add_paragraph('Use an in-browser, high-performance QR scanner to validate student tickets at the venue.', style='List Bullet')
doc.add_paragraph('Instantly mark attendance and verify ticket validity in real time.', style='List Bullet')
doc.add_paragraph('The system has built-in anti-fraud measures to prevent duplicate entries or invalid ticket usage.', style='List Bullet')

doc.add_heading('AI & Predictive Analytics Integration', level=1)
doc.add_paragraph('A standout feature of CollegeBuddy is its Machine Learning component. It uses a Linear Regression model to assist hosts in their logistical planning.')
doc.add_paragraph('Inputs: The model analyzes the event\'s participant limit, event fee, and total registrations.', style='List Bullet')
doc.add_paragraph('Output: It predicts the actual expected attendance.', style='List Bullet')
doc.add_paragraph('Goal: This helps organizers optimize resources like seating arrangements, catering, and venue size, reducing waste and saving money.', style='List Bullet')

doc.add_heading('Primary Use Cases', level=1)
doc.add_heading('1. Streamlining Event Discovery and Ticketing (For the Student Body)', level=2)
doc.add_paragraph('Currently, students might find out about events through scattered WhatsApp groups, emails, or bulletin boards. CollegeBuddy centralizes this. Students log in, view all happening activities, register frictionlessly, and get an organized digital QR ticket—eliminating the need for physical tickets or manual list-checking.')

doc.add_heading('2. Efficient Event Logistics & Resource Optimization (For College Administrations & Clubs)', level=2)
doc.add_paragraph('Organizing campus fests or workshops often results in over-ordering food or booking overly large halls because of unpredictable turnout. By using CollegeBuddy’s AI attendance prediction, organizers get a realistic number of expected attendees, allowing them to allocate budgets efficiently.')

doc.add_heading('3. Fraud-Proof Check-ins (For Event Security & Volunteers)', level=2)
doc.add_paragraph('Manual check-ins (e.g., using printed spreadsheets) are slow and prone to proxy entries. CollegeBuddy enables volunteers to stand at the gates with just their smartphones, scanning QR codes via the web scanner. The system instantly verifies if the ticket is valid and legitimate, checking the attendee in within seconds.')

doc.add_heading('4. Post-Event Analytics and Records (For Institutional Tracking)', level=2)
doc.add_paragraph('Colleges often need event data for accreditation or institutional reports. CollegeBuddy provides historical data, showing exactly how many students registered versus how many actually attended, generating actionable insights for future event planning.')

doc.add_heading('Future Scope & Roadmap', level=1)
doc.add_paragraph('Automated Certificates: Instantly generating and emailing participation certificates (PDFs) to users marked as "attended".', style='List Bullet')
doc.add_paragraph('Payment Gateway: Processing payments for high-value paid workshops directly on the platform.', style='List Bullet')
doc.add_paragraph('Push Notifications/Emails: Sending real-time reminders to students before an event starts.', style='List Bullet')
doc.add_paragraph('Multi-Campus Support: Scaling the system to support multiple institutions in a SaaS-like model.', style='List Bullet')

# Save in the same folder where the user is working
output_path = r'c:\Users\harsh\OneDrive\Desktop\CollegeBuddy\CollegeBuddy_Project_Information.docx'
doc.save(output_path)
print(f"Word document saved successfully at {output_path}")
